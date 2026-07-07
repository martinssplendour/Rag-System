"""Document loading: turn raw bytes/text into LangChain ``Document`` objects.

PDF pages are read with column-major coordinate sorting rather than
PyMuPDF's raw stream order or a simple row-band sort. Both dataset PDFs use
a strict left-column/right-column body layout (verified against the real
France/Italy files' block coordinates): the left column runs from the
title down to the final section, and the right column is a second,
independent stream of sections starting at the top of the page. A
row-band sort interleaves the two columns mid-sentence; the correct order
is every left-column block top-to-bottom, then every right-column block
top-to-bottom. See BUILD_SPEC_PART1_INGESTION.md section 1 (reference:
main spec section 9.3).
"""

from __future__ import annotations

from contextlib import redirect_stdout
from dataclasses import dataclass
from io import BytesIO, StringIO
from typing import Any

import fitz
from docx import Document as DocxDocument
from langchain_core.documents import Document

from app.rag.tables import StructuredTable


@dataclass(frozen=True)
class LoadedPdfPage:
    page_number: int
    text: str
    tables: list[StructuredTable]


def _column_major_sort_key(block: tuple, page_width: float) -> tuple[int, float]:
    x0, y0 = block[0], block[1]
    column = 0 if x0 < page_width / 2 else 1
    return (column, y0)


def _collapse_block_text(raw_block_text: str) -> str:
    """Join a block's internal word-wrap line breaks into a single line.

    PyMuPDF blocks contain "\\n" wherever the PDF layout engine wrapped a
    line, not wherever a real paragraph/heading boundary exists -- if left
    alone, a wrapped sentence fragment ending mid-clause can be mistaken
    for a section heading by split_into_sections. Collapsing here and
    joining blocks with a blank line (see load_pdf_pages) restores a
    paragraph-per-block structure that matches what the chunker expects.
    """
    return " ".join(raw_block_text.split())


def _normalise_table_cell(value: str | None) -> str:
    return " ".join(str(value or "").split())


def _rect_from_block(block: tuple) -> fitz.Rect:
    return fitz.Rect(float(block[0]), float(block[1]), float(block[2]), float(block[3]))


def _find_table_title(
    blocks: list[tuple],
    table_bbox: tuple[float, float, float, float],
) -> tuple[str | None, fitz.Rect | None]:
    table_rect = fitz.Rect(table_bbox)
    candidates: list[tuple[float, str, fitz.Rect]] = []
    for block in blocks:
        rect = _rect_from_block(block)
        if rect.y1 > table_rect.y0 + 2:
            continue
        if table_rect.y0 - rect.y1 > 90:
            continue
        if rect.x1 < table_rect.x0 or rect.x0 > table_rect.x1:
            continue
        text = _collapse_block_text(str(block[4]))
        if not text or len(text) > 100:
            continue
        candidates.append((rect.y1, text, rect))
    if not candidates:
        return None, None
    _, text, rect = max(candidates, key=lambda candidate: candidate[0])
    return text, rect


def _extract_tables_from_page(
    page: fitz.Page,
    *,
    page_number: int,
    blocks: list[tuple],
) -> tuple[list[StructuredTable], list[fitz.Rect]]:
    # PyMuPDF may print an optional package suggestion from find_tables().
    # Suppress that so ingestion logs only contain application messages.
    with redirect_stdout(StringIO()):
        finder = page.find_tables(strategy="lines")

    tables: list[StructuredTable] = []
    exclusion_rects: list[fitz.Rect] = []
    for table_index, table in enumerate(finder.tables, start=1):
        extracted_rows = table.extract()
        if len(extracted_rows) < 2:
            continue

        headers = [_normalise_table_cell(cell) for cell in extracted_rows[0]]
        if not any(headers):
            continue

        rows: list[dict[str, str]] = []
        for raw_row in extracted_rows[1:]:
            values = [_normalise_table_cell(cell) for cell in raw_row]
            if not any(values):
                continue
            rows.append(
                {
                    header: values[index] if index < len(values) else ""
                    for index, header in enumerate(headers)
                    if header
                }
            )
        if not rows:
            continue

        title, title_rect = _find_table_title(blocks, table.bbox)
        tables.append(
            StructuredTable(
                page_number=page_number,
                table_index=table_index,
                title=title,
                headers=headers,
                rows=rows,
                bounding_box=tuple(float(value) for value in table.bbox),
            )
        )
        exclusion_rects.append(fitz.Rect(table.bbox))
        if title_rect is not None:
            exclusion_rects.append(title_rect)

    return tables, exclusion_rects


def _block_is_excluded(block: tuple, exclusion_rects: list[fitz.Rect]) -> bool:
    rect = _rect_from_block(block)
    return any(rect.intersects(exclusion) for exclusion in exclusion_rects)


def load_pdf_pages_with_tables(pdf_bytes: bytes) -> list[LoadedPdfPage]:
    """Extract PDF prose and grid-line tables.

    Tables are extracted before prose text and their bounding boxes are
    removed from the block stream. This keeps table cells from being
    flattened into short fake headings during normal prose chunking.
    """
    document = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        pages: list[LoadedPdfPage] = []
        for zero_based_index, page in enumerate(document):
            page_number = zero_based_index + 1
            blocks = list(page.get_text("blocks"))
            tables, exclusion_rects = _extract_tables_from_page(
                page,
                page_number=page_number,
                blocks=blocks,
            )
            prose_blocks = [
                block for block in blocks if not _block_is_excluded(block, exclusion_rects)
            ]
            ordered = sorted(
                prose_blocks,
                key=lambda block: _column_major_sort_key(block, page.rect.width),
            )
            collapsed = [_collapse_block_text(block[4]) for block in ordered]
            text = "\n\n".join(block_text for block_text in collapsed if block_text)
            pages.append(LoadedPdfPage(page_number=page_number, text=text, tables=tables))
        return pages
    finally:
        document.close()


def load_pdf_pages(pdf_bytes: bytes) -> list[tuple[int, str]]:
    """Extract per-page prose text in reading order.

    Returns a list of ``(page_number, text)`` with 1-based page numbers.
    Grid-line table regions are excluded because they are represented as
    structured table chunks during ingestion.
    """
    return [(page.page_number, page.text) for page in load_pdf_pages_with_tables(pdf_bytes)]


def load_pdf(pdf_bytes: bytes, metadata: dict[str, Any]) -> list[Document]:
    return [
        Document(page_content=text, metadata={**metadata, "page_number": page_number})
        for page_number, text in load_pdf_pages(pdf_bytes)
    ]


def load_docx_text(docx_bytes: bytes) -> str:
    document = DocxDocument(BytesIO(docx_bytes))
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    return "\n\n".join(blocks)


def load_docx(docx_bytes: bytes, metadata: dict[str, Any]) -> list[Document]:
    return [Document(page_content=load_docx_text(docx_bytes), metadata=metadata)]


def load_text(text: str, metadata: dict[str, Any]) -> list[Document]:
    return [Document(page_content=text, metadata=metadata)]
